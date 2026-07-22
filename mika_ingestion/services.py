"""
services.py
-----------
Thin wrappers around the RAG modules (builder.py, git_builder.py, etc.).
All paths are resolved from Django settings so the rest of the app stays clean.

Supported ingestion sources:
  • URL    – crawled with Firecrawl (ingest_url)
  • GitHub – repo clone/pull       (ingest_github)
  • File   – PDF, Markdown, DOCX   (ingest_file)  ← replaces ingest_pdf

Settings required in mika_project/settings.py:
  RAG_DIR            – folder with chunker.py, embedder.py, etc.
  INDEX_DIR          – where .faiss + .pkl pairs live
  UPLOAD_DIR         – temp scratch folder for uploads
  MAX_UPLOAD_BYTES   – maximum accepted file size
"""

import sys
import os
import logging
from pathlib import Path

from django.conf import settings

logger = logging.getLogger(__name__)

# ── Extensions we accept for file uploads ────────────────────────────────────
SUPPORTED_EXTENSIONS = {".pdf", ".md", ".markdown", ".docx", ".doc"}


def _ensure_rag_on_path():
    """Add RAG_DIR to sys.path so we can import the original scripts."""
    rag = str(settings.RAG_DIR)
    if rag not in sys.path:
        sys.path.insert(0, rag)


def _index_paths(index_name: str):
    """Return (faiss_path, pkl_path) for a given index name."""
    base = settings.INDEX_DIR / index_name
    return str(base) + ".faiss", str(base) + ".pkl"


def list_indexes():
    """Return a list of existing index names (stems of .faiss files)."""
    return sorted(p.stem for p in settings.INDEX_DIR.glob("*.faiss"))


# ── Text extractors ───────────────────────────────────────────────────────────

def _extract_pdf(path: str) -> str:
    """Extract plain text from a PDF using PyMuPDF (pip install pymupdf)."""
    import fitz  # PyMuPDF
    doc  = fitz.open(path)
    text = "\n".join(page.get_text() for page in doc)
    doc.close()
    logger.info("PDF extracted: %d chars from %s", len(text), path)
    return text


def _extract_markdown(path: str) -> str:
    """Read a Markdown file as plain text (the raw markdown is fine for RAG)."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    logger.info("Markdown read: %d chars from %s", len(text), path)
    return text


def _extract_docx(path: str) -> str:
    """
    Extract text from a .docx file using python-docx.
    Install with: pip install python-docx
    """
    from docx import Document  # python-docx
    doc   = Document(path)
    lines = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also grab text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    text = "\n".join(lines)
    logger.info("DOCX extracted: %d chars from %s", len(text), path)
    return text


def _extract_text(file_path: str) -> str:
    """Dispatch to the right extractor based on file extension."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(file_path)
    if ext in (".md", ".markdown"):
        return _extract_markdown(file_path)
    if ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    raise ValueError(
        f"Unsupported file type '{ext}'. "
        f"Accepted: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


# ── Shared index builder ──────────────────────────────────────────────────────

def _normalize_chunks(chunks: list) -> list[dict]:
    """
    Normalize any mix of bare strings and dicts into uniform
    {"text": str, "source": str} dicts.

    Indexes built before the dict-chunk migration store plain strings.
    Mixing old strings + new dicts in the same pkl causes subtle bugs
    (KeyError on retrieval, inaccurate source attribution). This function
    makes both formats safe to merge.
    """
    out = []
    for c in chunks:
        if isinstance(c, dict) and "text" in c:
            out.append(c)
        else:
            out.append({"text": str(c), "source": ""})
    return out


def _build_or_append(new_chunks: list[dict], index_name: str, append: bool) -> dict:
    """
    Given a list of chunk dicts {"text": ..., "source": ...},
    build a fresh index or append to an existing one.
    Returns {"chunks": int, "vectors": int, "index_name": str}.
    """
    _ensure_rag_on_path()
    import faiss
    import pickle
    from embedder import Embedder
    from indexer  import Indexer
    from storage  import Storage

    if not new_chunks:
        raise ValueError(
            "No chunks were produced from the source. "
            "The content may be empty, all-whitespace, or the file/URL returned no usable text."
        )

    embedder = Embedder()
    storage  = Storage()

    # Ensure new chunks are always dicts (defensive normalisation)
    new_chunks = _normalize_chunks(new_chunks)

    # Extract text strings for embedding (embedder expects list[str])
    texts          = [c["text"] for c in new_chunks]
    new_embeddings = embedder.embed(texts)

    index_path, chunks_path = _index_paths(index_name)

    if append and Path(index_path).exists():
        logger.info("Appending to existing index '%s'", index_name)
        old_index = faiss.read_index(index_path)
        with open(chunks_path, "rb") as f:
            raw_old = pickle.load(f)

        # ← KEY FIX: normalize old chunks so string + dict mixing never happens
        old_chunks = _normalize_chunks(raw_old)

        indexer = Indexer(old_index.d)
        indexer.index = old_index
        merged_index  = indexer.append_to_index(new_embeddings)
        merged_chunks = old_chunks + new_chunks

        logger.info(
            "Merged %d old + %d new = %d total chunks",
            len(old_chunks), len(new_chunks), len(merged_chunks),
        )
    else:
        logger.info("Building new index '%s'", index_name)
        dim           = new_embeddings.shape[1]
        indexer       = Indexer(dim)
        merged_index  = indexer.build_index(new_embeddings)
        merged_chunks = new_chunks

    storage.save_index(merged_index, index_path)
    storage.save_chunks(merged_chunks, chunks_path)

    # Explicitly bump the mtime on the .faiss file so that any running bot
    # process detects the change via its mtime-based hot-reload and picks up
    # the new data without needing a server restart.
    Path(index_path).touch()

    logger.info(
        "Index '%s': %d chunks, %d vectors",
        index_name, len(merged_chunks), merged_index.ntotal,
    )
    return {
        "chunks":     len(merged_chunks),
        "vectors":    merged_index.ntotal,
        "index_name": index_name,
    }


# ── Public ingestion API ──────────────────────────────────────────────────────

def ingest_file(file_path: str, index_name: str, append: bool,
                original_name: str = "") -> dict:
    """
    Extract text from a PDF, Markdown, or DOCX file, chunk it,
    embed it, and build/update a FAISS index.

    Cleans up the temp file when done.
    """
    _ensure_rag_on_path()
    from chunker import Chunker

    source   = original_name or Path(file_path).name
    raw_text = _extract_text(file_path)

    if not raw_text.strip():
        raise ValueError(
            f"No readable text found in '{source}'. "
            "If it's a scanned PDF, OCR is not yet supported."
        )

    chunker    = Chunker()
    new_chunks = chunker.chunk_text(raw_text, source=source)

    try:
        return _build_or_append(new_chunks, index_name, append)
    finally:
        try:
            os.remove(file_path)
        except OSError:
            pass


# Keep old name as alias so any code still calling ingest_pdf() doesn't break
def ingest_pdf(pdf_path: str, index_name: str, append: bool) -> dict:
    return ingest_file(pdf_path, index_name, append, original_name=Path(pdf_path).name)


def ingest_url(url: str, index_name: str, append: bool) -> dict:
    """Crawl a website with Firecrawl and build/update a FAISS index."""
    _ensure_rag_on_path()
    from chunker import Chunker
    from crawler import Crawler

    logger.info("Crawling %s …", url)
    crawler    = Crawler()
    raw_docs   = crawler.crawler_job(url)

    chunker    = Chunker()
    new_chunks = chunker.chunk_docs(raw_docs, source=url)

    return _build_or_append(new_chunks, index_name, append)


def ingest_github(repo_url: str, index_name: str, append: bool,
                  extensions: list[str] | None = None) -> dict:
    """Clone/pull a GitHub repo and build/update a FAISS index."""
    _ensure_rag_on_path()
    from chunker      import Chunker
    from git_builder  import GitHubBuilder

    # Strip trailing .git and slashes so the cache dir name is stable
    # e.g. https://github.com/user/repo.git → repo
    clean_url = repo_url.rstrip("/")
    repo_slug = clean_url.split("/")[-1].removesuffix(".git")
    repo_dir  = str(settings.INDEX_DIR / f"_repo_{repo_slug}")

    builder = GitHubBuilder(repo_url=repo_url, repo_dir=repo_dir)

    # clone_repo() must be called before load_files() so the directory exists
    builder.clone_repo()

    exts     = extensions or [".py", ".js", ".ts", ".md"]
    raw_docs = builder.load_files(exts=exts)

    if not raw_docs:
        tried = ", ".join(sorted(set(exts)))
        raise ValueError(
            f"No files with extensions [{tried}] found in repo '{repo_slug}'. "
            f"Check that the repo contains files with those extensions, "
            f"or select different extensions in the UI."
        )

    chunker    = Chunker()
    new_chunks = chunker.chunk_docs(raw_docs, source=repo_url)

    return _build_or_append(new_chunks, index_name, append)
